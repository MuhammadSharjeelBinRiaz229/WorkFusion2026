import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Paths
script_dir = os.path.dirname(os.path.abspath(__file__))
base_dir = os.path.dirname(script_dir)
docs_dir = os.path.join(os.path.dirname(base_dir), "docs")
md_path = os.path.join(docs_dir, "AI_Recommendation_Report.md")
docx_path = os.path.join(docs_dir, "AI_Recommendation_Report.docx")

def set_cell_background(cell, fill_hex):
    """Applies solid background color to table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 dxa = 1/20th of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_formatted_text(paragraph, text, default_font="Calibri", default_size=11, color_rgb=(51, 51, 51)):
    """Parses text for bold (**text**) and code (`code`) inline structures and adds them to paragraph."""
    # Pattern to match bold or inline code
    pattern = re.compile(r'(\*\*.*?\*\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
            
        run = paragraph.add_run()
        run.font.name = default_font
        run.font.size = Pt(default_size)
        run.font.color.rgb = RGBColor(*color_rgb)
        
        if part.startswith('**') and part.endswith('**'):
            # Bold run
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code run
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(default_size - 1)
            run.font.color.rgb = RGBColor(180, 40, 40)
        else:
            run.text = part

def build_docx():
    print(f"Reading Markdown from: {md_path}")
    if not os.path.exists(md_path):
        print("Error: Markdown report does not exist.")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # Configure page margins (1 inch = 72 pt)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Title Block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("WorkFusion AI Recommendation Engine\nTechnical Report")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(26)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102) # Deep blue

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(24)
    meta_run = meta_p.add_run("Document Version: 1.0  |  Status: Production Ready  |  Pakistan's Hybrid Marketplace")
    meta_run.font.name = "Calibri"
    meta_run.font.size = Pt(10.5)
    meta_run.italic = True
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    # State variables for parsing
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []
    
    for idx, line in enumerate(lines):
        clean_line = line.strip()
        
        # Skip the title and meta blocks in markdown as we created them manually above
        if idx < 6 and (clean_line.startswith("# WorkFusion") or "Document Version:" in clean_line):
            continue
            
        # Parse Code Block
        if clean_line.startswith("```"):
            if in_code_block:
                # Close code block
                in_code_block = False
                code_text = "\n".join(code_block_lines)
                
                # Add code block with gray background
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Add background shading via XML
                pBrd = OxmlElement('w:pBdr')
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '24') # 3pt
                left_border.set(qn('w:space'), '4')
                left_border.set(qn('w:color'), 'A0A0A0')
                pBrd.append(left_border)
                p._p.get_or_add_pPr().append(pBrd)
                
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                p._p.get_or_add_pPr().append(shd)
                
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
                code_block_lines = []
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            # Simply accumulate lines
            code_block_lines.append(line.replace('\n', ''))
            continue

        # Parse Table Blocks
        if clean_line.startswith("|") and not in_code_block:
            if not in_table:
                in_table = True
            table_lines.append(clean_line)
            continue
        elif in_table:
            # End of table detected
            in_table = False
            
            # Process table_lines
            # Filter out lines that are separator lines (e.g., contains only dashes, colons, pipes)
            filtered_rows = []
            for t_line in table_lines:
                # check if separator
                if re.match(r'^\|[\s:-|]+$', t_line):
                    continue
                # Split cell values
                cells = [c.strip() for c in t_line.split('|')]
                if len(cells) > 1:
                    # Remove empty first/last elements from split
                    if cells[0] == '': cells.pop(0)
                    if cells and cells[-1] == '': cells.pop()
                    filtered_rows.append(cells)
            
            if filtered_rows:
                num_cols = len(filtered_rows[0])
                num_rows = len(filtered_rows)
                
                docx_table = doc.add_table(rows=num_rows, cols=num_cols)
                docx_table.style = 'Light Shading Accent 1'
                
                # Popualte cells
                for r_idx, row_data in enumerate(filtered_rows):
                    row = docx_table.rows[r_idx]
                    for c_idx, val in enumerate(row_data):
                        cell = row.cells[c_idx]
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)
                        
                        # Add value
                        add_formatted_text(p, val, default_size=10, color_rgb=(0,0,0) if r_idx==0 else (51,51,51))
                        
                        # Style table headers
                        if r_idx == 0:
                            set_cell_background(cell, "003366") # dark blue header
                            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # Make header run bold
                            for run in p.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255) # white text
                        else:
                            # Soft light row backgrounds
                            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                            if r_idx % 2 == 0:
                                set_cell_background(cell, "F9FBFD")
                            else:
                                set_cell_background(cell, "FFFFFF")
                                
            table_lines = []
            # Do not continue, parse current line as a normal line below
            
        # Parse Headings
        if clean_line.startswith("#"):
            h_match = re.match(r'^(#+)\s+(.*)$', clean_line)
            if h_match:
                level = len(h_match.group(1))
                h_text = h_match.group(2)
                
                # Add heading (1-indexed for levels)
                heading = doc.add_heading(level=level)
                heading.paragraph_format.space_before = Pt(16)
                heading.paragraph_format.space_after = Pt(6)
                heading.paragraph_format.keep_with_next = True
                
                run = heading.add_run(h_text)
                run.font.name = "Calibri"
                run.bold = True
                
                if level == 1:
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(0, 51, 102) # Deep blue
                elif level == 2:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 102, 204) # Medium blue
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(51, 51, 51) # Dark gray
                continue
                
        # Parse Bullet Lists
        if clean_line.startswith("- ") or clean_line.startswith("* "):
            bullet_text = clean_line[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            add_formatted_text(p, bullet_text)
            continue
            
        # Parse Images: ![alt](path)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', clean_line)
        if img_match:
            img_alt = img_match.group(1)
            img_rel_path = img_match.group(2)
            
            # Resolve image path relative to docs folder
            full_img_path = os.path.join(docs_dir, img_rel_path)
            if os.path.exists(full_img_path):
                # Add image paragraph
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(12)
                img_p.paragraph_format.space_after = Pt(4)
                
                run = img_p.add_run()
                run.add_picture(full_img_path, width=Inches(5.8))
                
                # Add caption
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_after = Pt(12)
                cap_run = cap_p.add_run(f"Figure: {img_alt}")
                cap_run.font.name = "Calibri"
                cap_run.font.size = Pt(9.5)
                cap_run.italic = True
                cap_run.font.color.rgb = RGBColor(120, 120, 120)
            else:
                print(f"Warning: Image file not found at {full_img_path}")
            continue

        # Horizontal rules
        if clean_line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            # Draw horizontal border line XML
            pBrd = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6') # thin line
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), 'D3D3D3')
            pBrd.append(bottom_border)
            p._p.get_or_add_pPr().append(pBrd)
            continue
            
        # Standard Paragraphs
        if clean_line:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, clean_line)
            
    print(f"Saving Word Document to: {docx_path}")
    doc.save(docx_path)
    print("Document build completed successfully!")

if __name__ == "__main__":
    build_docx()
